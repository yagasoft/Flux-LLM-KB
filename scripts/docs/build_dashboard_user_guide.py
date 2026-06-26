from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


REPO_ROOT = Path(__file__).resolve().parents[2]
GUIDE_DIR = REPO_ROOT / "docs" / "user-guide"
SCREENS_DIR = GUIDE_DIR / "screens"
MANUAL_MD = GUIDE_DIR / "dashboard-user-manual.md"
DOCX_PATH = GUIDE_DIR / "Flux-LLM-KB-Dashboard-User-Manual.docx"


@dataclass(frozen=True)
class ScreenSpec:
    slug: str
    title: str
    subtitle: str
    active_nav: str
    cards: tuple[tuple[str, str, str], ...]
    panels: tuple[tuple[str, tuple[str, ...]], ...]


SCREEN_SPECS: tuple[ScreenSpec, ...] = (
    ScreenSpec(
        slug="overview",
        title="Overview",
        subtitle="Read-only system status and next safe action",
        active_nav="Overview",
        cards=(
            ("System", "Ready", "Postgres, API, and dashboard responding"),
            ("Attention", "2 items", "One stale evidence gate, one blocked job"),
            ("Flux handled", "5 actions", "Evidence refreshed and embeddings queued"),
            ("Next safe action", "Run guarded pass", "No settings mutation required"),
        ),
        panels=(
            ("What needs attention", ("Indexer evidence is 18 hours old", "One Office job is blocked by a missing dependency")),
            ("Flux handled automatically", ("Queued 8 embedding refreshes", "Ran governance shadow proposals", "Cleared completed diagnostic errors")),
            ("Recommended safe action", ("Open Automation and run a guarded pass now", "Review blocked items before manual work")),
        ),
    ),
    ScreenSpec(
        slug="automation",
        title="Automation",
        subtitle="Guarded Auto posture with durable action history",
        active_nav="Automation",
        cards=(
            ("Posture", "Guarded Auto", "Recurring setting is disabled by default"),
            ("Last run", "10:42", "4 safe actions, 0 setting changes"),
            ("Next run", "Manual", "Operator can run a guarded pass"),
            ("Settings mutated", "false", "Guarded pass cannot change runtime settings"),
        ),
        panels=(
            ("Eligible actions", ("Refresh retrieval evidence", "Ingest approved captures", "Enqueue missing embeddings", "Run governance shadow proposals")),
            ("Manual required", ("Deletes and purge actions", "OAuth and host startup", "Restart or reindex settings", "Open or reveal local files")),
            ("Automation audit trail", ("completed - refresh_retrieval_evidence", "completed - enqueue_embedding_refresh", "manual_required - restart_settings")),
        ),
    ),
    ScreenSpec(
        slug="diagnostics",
        title="Diagnostics",
        subtitle="Structured errors, filters, details, remediation, and navigation",
        active_nav="Diagnostics",
        cards=(
            ("Open errors", "3", "Filtered to actionable items"),
            ("Safe fixes", "2", "Retry job and clear completed errors"),
            ("Blocked", "1", "Missing dependency remains manual"),
            ("Settings mutated", "false", "Remediation actions stay bounded"),
        ),
        panels=(
            ("Filters", ("Section: all", "Status: blocked_missing_dependency", "Family: office")),
            ("Structured error detail", ("Target: job example-office-12", "Reason: extractor dependency missing", "Copy detail or navigate to Jobs")),
            ("Remediation", ("Retry corpus job", "Repair asset statuses", "Clear completed errors")),
        ),
    ),
    ScreenSpec(
        slug="performance",
        title="Performance",
        subtitle="Acceleration capability, reliability gates, and worker telemetry",
        active_nav="Performance",
        cards=(
            ("CPU", "Ready", "4 worker families available"),
            ("Cache", "Warm", "ASR and extraction cache reachable"),
            ("Reliability", "Pass", "Latest root gate succeeded"),
            ("Benchmarks", "12 runs", "Synthetic and monitored-root history"),
        ),
        panels=(
            ("Acceleration capability", ("CPU and disk hints are healthy", "Optional ONNX provider unavailable", "Local model probe disabled")),
            ("Reliability gates", ("All-root reliability evidence fresh", "Watcher probe passed", "No slow lock backlog")),
            ("Worker telemetry", ("office: 1 queued", "media: 0 blocked", "embeddings: 8 queued")),
        ),
    ),
    ScreenSpec(
        slug="corpus",
        title="Corpus",
        subtitle="Watched roots, sync policy, and crawl jobs",
        active_nav="Corpus",
        cards=(
            ("Roots", "3", "Docs, Projects, Knowledge"),
            ("Watching", "2", "Native watcher active"),
            ("Backfill", "Ready", "Embeddings and extraction available"),
            ("Policy", "Inherited", "Root globs extend global defaults"),
        ),
        panels=(
            ("Watched roots", ("Docs root - enabled", "Projects root - polling fallback", "Archive root - disabled")),
            ("Safe actions", ("Sync selected root", "Dry-run a scoped path", "Run bounded backfill")),
            ("Manual actions", ("Delete root", "Purge assets", "Broad reindex settings")),
        ),
    ),
    ScreenSpec(
        slug="mail",
        title="Mail",
        subtitle="Profiles, OAuth status, sync, and post-process policy",
        active_nav="Mail",
        cards=(
            ("Profiles", "2", "Gmail capture and Outlook catch-up"),
            ("OAuth", "Needs review", "Gmail profile token expires soon"),
            ("Sync", "Idle", "Manual sync available"),
            ("Approved captures", "4", "Eligible for guarded ingestion"),
        ),
        panels=(
            ("Profile inspector", ("gmail-capture - IMAP", "outlook-catchup - Outlook host", "Post-process dry-run available")),
            ("Manual controls", ("Start Gmail OAuth", "Start host process", "Change destructive mail policy")),
            ("Guarded automation", ("Ingest already-approved captures only", "Record sanitized job IDs and counts")),
        ),
    ),
    ScreenSpec(
        slug="retrieval",
        title="Retrieval",
        subtitle="Search quality, explain traces, benchmarks, and code diagnostics",
        active_nav="Retrieval",
        cards=(
            ("Search", "Ready", "Local-first retrieval active"),
            ("Explain", "Enabled", "Ranking and suppression shown"),
            ("Benchmarks", "Fresh", "Standard suite completed"),
            ("Code diagnostics", "Moved here", "Coverage belongs with retrieval"),
        ),
        panels=(
            ("Retrieval controls", ("Search with filters", "Explain confidence bands", "Build compact brief")),
            ("Code diagnostics", ("Parser coverage summary", "Fallback and generated-file handling", "Sanitized miss feedback")),
            ("Benchmark history", ("standard - pass", "governance-shadow - advisory", "Calibration delta visible")),
        ),
    ),
    ScreenSpec(
        slug="review",
        title="Review",
        subtitle="Human judgment for captures, claims, and governance",
        active_nav="Review",
        cards=(
            ("Capture queue", "6", "Rationale required"),
            ("Governance", "Shadow", "Proposals only"),
            ("Duplicates", "4 clusters", "Advisory review"),
            ("Recoverable", "2 actions", "Before-state captured"),
        ),
        panels=(
            ("Capture review", ("Approve with rationale", "Reject with rationale", "Ingest after approval")),
            ("Governance proposals", ("Shadow proposal: merge duplicate claim", "Manual risk: apply required", "Recovery action available")),
            ("Memory quality", ("Needs review", "Suppressed duplicate", "Retention candidate")),
        ),
    ),
    ScreenSpec(
        slug="settings",
        title="Settings",
        subtitle="System settings, Codex hooks, deployment, runtime, restart, and reindex",
        active_nav="Settings",
        cards=(
            ("Codex hooks", "Ready", "kb.brief and finalize hooks active"),
            ("Deployment", "Local", "API and dashboard reachable"),
            ("Runtime actions", "2", "Acknowledged and pending requests"),
            ("Automation", "Disabled", "operator.automation.enabled default false"),
        ),
        panels=(
            ("Runtime settings", ("retrieval.token_budget - live", "embedding.dimensions - reindex_required", "operator.automation.mode - guarded")),
            ("System sections", ("Codex Hooks", "Deployment", "Runtime Actions", "Restart / Reindex Changes")),
            ("Manual controls", ("Apply acknowledged restart request", "Confirm reindex-class changes", "Reset setting to catalog default")),
        ),
    ),
    ScreenSpec(
        slug="jobs",
        title="Jobs",
        subtitle="Background work by status, family, and age",
        active_nav="Jobs",
        cards=(
            ("Queued", "9", "Mostly embeddings"),
            ("Running", "2", "Office extraction and crawl sync"),
            ("Blocked", "1", "Missing dependency"),
            ("Completed", "41", "Last 24 hours"),
        ),
        panels=(
            ("Worker families", ("office - 1 running", "media - idle", "embeddings - 8 queued", "mail - idle")),
            ("Job list", ("example-office-12 - blocked_missing_dependency", "example-embed-08 - queued", "example-sync-21 - completed")),
            ("Next action", ("Use Diagnostics for safe retries", "Use Performance for worker telemetry")),
        ),
    ),
    ScreenSpec(
        slug="result-detail",
        title="Result Details",
        subtitle="Safe detail drawer for search and file results",
        active_nav="Retrieval",
        cards=(
            ("Result", "File evidence", "Preview available"),
            ("Explanation", "Visible", "Ranking and filter trace shown"),
            ("Copy", "Allowed", "Safe metadata copied"),
            ("Open file", "Manual", "Reveal requires explicit operator action"),
        ),
        panels=(
            ("Preview", ("Public-safe excerpt placeholder", "Chunk metadata and score", "No raw private path in guide screenshot")),
            ("Actions", ("Copy detail", "Go to owning tab", "Manual open/reveal when allowed")),
            ("Why this result", ("Local workspace boost", "Current-state filter", "Suppression trace clear")),
        ),
    ),
)

NAV_ITEMS = ("Overview", "Automation", "Diagnostics", "Performance", "Corpus", "Mail", "Settings", "Retrieval", "Review", "Jobs")


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = (
        Path("C:/Windows/Fonts/segoeuib.ttf" if bold else "C:/Windows/Fonts/segoeui.ttf"),
        Path("C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf"),
        Path("/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf" if bold else "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"),
    )
    for candidate in candidates:
        if candidate.exists():
            return ImageFont.truetype(str(candidate), size=size)
    return ImageFont.load_default()


def _wrap(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.ImageFont, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for word in text.split():
        trial = f"{current} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= width:
            current = trial
            continue
        if current:
            lines.append(current)
        current = word
    if current:
        lines.append(current)
    return lines


def _text(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font: ImageFont.ImageFont, fill: str = "#172033", max_width: int | None = None, line_gap: int = 6) -> int:
    x, y = xy
    lines = [text] if max_width is None else _wrap(draw, text, font, max_width)
    for line in lines:
        draw.text((x, y), line, fill=fill, font=font)
        bbox = draw.textbbox((x, y), line, font=font)
        y += bbox[3] - bbox[1] + line_gap
    return y


def _rounded(draw: ImageDraw.ImageDraw, box: tuple[int, int, int, int], fill: str, outline: str = "#d9e2ef", radius: int = 10) -> None:
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=1)


def capture_screens() -> None:
    SCREENS_DIR.mkdir(parents=True, exist_ok=True)
    for spec in SCREEN_SPECS:
        image = Image.new("RGB", (1600, 940), "#f4f7fb")
        draw = ImageDraw.Draw(image)
        title_font = _font(40, True)
        h2_font = _font(24, True)
        h3_font = _font(18, True)
        body_font = _font(17)
        small_font = _font(14)
        label_font = _font(13, True)

        draw.rectangle((0, 0, 250, 940), fill="#172033")
        _text(draw, (34, 32), "FLUX", _font(30, True), "#ffffff")
        _text(draw, (34, 68), "LLM-KB", small_font, "#a8b6c9")
        nav_y = 126
        for item in NAV_ITEMS:
            active = item == spec.active_nav
            if active:
                draw.rounded_rectangle((22, nav_y - 8, 228, nav_y + 36), radius=9, fill="#e8f4ff")
            _text(draw, (42, nav_y), item, body_font, "#172033" if active else "#d5deea")
            nav_y += 54

        draw.rectangle((250, 0, 1600, 92), fill="#ffffff")
        _text(draw, (292, 26), spec.title, title_font, "#172033")
        _text(draw, (292, 67), spec.subtitle, small_font, "#52627a")
        _rounded(draw, (1250, 26, 1518, 66), "#f9fbfe")
        _text(draw, (1270, 38), "Search memories, mail, corpus...", small_font, "#7a8799")

        card_colors = ("#e9f8f0", "#eaf2ff", "#fff4e6", "#f1edff")
        card_x = 292
        for index, (label, value, desc) in enumerate(spec.cards):
            box = (card_x + index * 315, 124, card_x + index * 315 + 285, 238)
            _rounded(draw, box, card_colors[index % len(card_colors)], outline="#d6e4f2")
            _text(draw, (box[0] + 18, box[1] + 18), label.upper(), label_font, "#52627a")
            _text(draw, (box[0] + 18, box[1] + 44), value, h2_font, "#172033")
            _text(draw, (box[0] + 18, box[1] + 78), desc, small_font, "#52627a", max_width=245)

        panel_positions = ((292, 282, 694, 770), (720, 282, 1122, 770), (1148, 282, 1550, 770))
        for panel, box in zip(spec.panels, panel_positions):
            heading, bullets = panel
            _rounded(draw, box, "#ffffff", outline="#d9e2ef")
            _text(draw, (box[0] + 22, box[1] + 20), heading, h2_font, "#172033")
            y = box[1] + 72
            for bullet in bullets:
                draw.ellipse((box[0] + 24, y + 7, box[0] + 34, y + 17), fill="#2f855a")
                y = _text(draw, (box[0] + 48, y), bullet, body_font, "#2c3748", max_width=box[2] - box[0] - 74, line_gap=8) + 10

        _rounded(draw, (292, 805, 1550, 888), "#ffffff", outline="#d9e2ef")
        _text(draw, (318, 828), "Public-safe mocked screenshot", h3_font, "#172033")
        _text(
            draw,
            (318, 858),
            "All names, counts, roots, accounts, and jobs are sample data for the user manual.",
            small_font,
            "#52627a",
        )
        image.save(SCREENS_DIR / f"{spec.slug}.png", optimize=True)


def _set_doc_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9.5)

    for style_name, size, color in (
        ("Heading 1", 18, "172033"),
        ("Heading 2", 14, "172033"),
        ("Heading 3", 11, "172033"),
    ):
        style = document.styles[style_name]
        style.font.name = "Aptos"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def _clean_inline(text: str) -> str:
    return text.replace("`", "")


def _add_image(document: Document, alt: str, rel_path: str) -> None:
    image_path = (GUIDE_DIR / rel_path).resolve()
    if not image_path.exists():
        raise FileNotFoundError(f"Referenced screenshot not found: {image_path}")
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(6.8))
    caption = document.add_paragraph(alt)
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.runs[0].font.size = Pt(8)
    caption.runs[0].font.italic = True
    caption.runs[0].font.color.rgb = RGBColor(82, 98, 122)


def build_docx() -> None:
    if not MANUAL_MD.exists():
        raise FileNotFoundError(MANUAL_MD)
    document = Document()
    _set_doc_styles(document)
    document.core_properties.title = "Flux LLM KB Dashboard User Manual"
    document.core_properties.subject = "Dashboard, guarded automation, diagnostics, performance, and operator workflows"
    document.core_properties.comments = "Generated from docs/user-guide/dashboard-user-manual.md using compact_reference_guide layout."

    first_h2 = True
    in_code = False
    for raw_line in MANUAL_MD.read_text(encoding="utf-8").splitlines():
        line = raw_line.rstrip()
        if line.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            paragraph = document.add_paragraph(line)
            paragraph.style = document.styles["Normal"]
            for run in paragraph.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(8.5)
            continue
        if not line:
            continue
        image_match = re.match(r"!\[(?P<alt>[^\]]+)\]\((?P<path>[^)]+)\)", line)
        if image_match:
            _add_image(document, image_match.group("alt"), image_match.group("path"))
            continue
        if line.startswith("# "):
            title = document.add_heading(line[2:], level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if line.startswith("## "):
            if first_h2:
                first_h2 = False
            else:
                document.add_page_break()
            document.add_heading(line[3:], level=1)
            continue
        if line.startswith("### "):
            document.add_heading(line[4:], level=2)
            continue
        if line.startswith("- "):
            document.add_paragraph(_clean_inline(line[2:]), style="List Bullet")
            continue
        paragraph = document.add_paragraph(_clean_inline(line))
        paragraph.paragraph_format.space_after = Pt(4)

    DOCX_PATH.parent.mkdir(parents=True, exist_ok=True)
    document.save(DOCX_PATH)


def main() -> int:
    parser = argparse.ArgumentParser(description="Build the Flux dashboard user guide screenshots and DOCX.")
    parser.add_argument("--capture-screens", action="store_true", help="Generate public-safe mocked dashboard screenshots.")
    parser.add_argument("--build-docx", action="store_true", help="Build the DOCX from markdown and screenshots.")
    parser.add_argument("--all", action="store_true", help="Generate screenshots and build the DOCX.")
    args = parser.parse_args()

    if not (args.capture_screens or args.build_docx or args.all):
        args.all = True
    if args.capture_screens or args.all:
        capture_screens()
    if args.build_docx or args.all:
        build_docx()
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
